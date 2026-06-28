from datetime import datetime, timedelta
import os
import json
from dotenv import load_dotenv

import lxml
from lxml import etree as LET

import docx
from docx.oxml import OxmlElement
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import RGBColor
from diff_match_patch import diff_match_patch

import zipfile
import xml.etree.ElementTree as ET

from interface_to_LLM.interface_to_llm import InterfaceToLLM, AnthropicClient

import re

# ---------------------------------------------------------------------------
# CHANGE vs v1: Section batch size — how many paragraphs to send per LLM call.
# Larger = better local coherence, more tokens per call.
# Smaller = faster, easier to debug. 10–20 is a good starting range.
# ---------------------------------------------------------------------------
BATCH_SIZE = 15
USE_RELEVANT_TERMS = False  # Set to True to inject per-batch terminology matches into prompts

# ---------------------------------------------------------------------------
# FORCED_CORRECTIONS: deterministic substitutions for terms the model
# consistently miscapitalises due to ambiguity with proper noun forms.
# "Subsidiary Body" is a proper noun when naming SBI/SBSTA specifically,
# but "subsidiary bodies" (generic plural) must be lower-cased.
# Add any further terms here as needed.
# ---------------------------------------------------------------------------
FORCED_CORRECTIONS = {
    "Subsidiary Bodies": "subsidiary bodies",
    "the Subsidiary Body": "the subsidiary body",
    # "Global Goal on Adaptation" is not a proper noun — lowercase per style guide
    "Global Goal on Adaptation": "global goal on adaptation", #"Global goal on adaptation" if first word in a sentence
    "Global Goal on adaptation": "global goal on adaptation", #"Global goal on adaptation" if first word in a sentence
    # CTC and CTCN are distinct — do not substitute one for the other
    # (handled via system prompt but also guarded here as a safety net)
    # "Party" is always capitalised in UNFCCC context — including compound modifiers.
    # The model treats "party-driven" etc. as generic lowercase; force correct form.
    "party-driven": "Party-driven",
    "party-led": "Party-led",
    "party-centric": "Party-centric",
    "non-party": "non-Party",
    # Note: "Annex" vs "annex" is context-dependent and cannot be handled here.
    # "Annex" (capital) = annex to the Convention or Kyoto Protocol.
    # "annex" (lower) = annex to any other document.
    # This is handled via the system prompt and pre-pass instructions.
}

def apply_forced_corrections(text: str) -> str:
    """Applies deterministic corrections for terms the model consistently gets wrong."""
    for wrong, correct in FORCED_CORRECTIONS.items():
        text = text.replace(wrong, correct)
    return text


# Acronyms that must never be substituted for another acronym.
# If the model replaces any of these with a different acronym, it is reversed.
PROTECTED_ACRONYMS = {"CTC", "CTCN", "COP", "SBI", "SBSTA", "GCF", "GEF", "CDM", "LEG", "SCF"}

def protect_acronyms(edited_text: str, original_text: str) -> str:
    """
    Reverses substitutions where the model replaced one protected acronym with
    a different protected acronym (e.g. CTC → CTCN).
    """
    import re as _re
    result = edited_text

    # --- Guard 1: reverse acronym-for-acronym substitutions ---
    for orig_acr in PROTECTED_ACRONYMS:
        orig_count = original_text.count(orig_acr)
        if orig_count == 0:
            continue
        edit_count = result.count(orig_acr)
        if edit_count >= orig_count:
            continue
        for sub_acr in PROTECTED_ACRONYMS:
            if sub_acr == orig_acr:
                continue
            sub_count_orig = original_text.count(sub_acr)
            sub_count_edit = result.count(sub_acr)
            excess = sub_count_edit - sub_count_orig
            if excess <= 0:
                continue
            if edit_count > 0:
                continue
            replaced = 0
            pos = 0
            while replaced < excess:
                idx = result.find(sub_acr, pos)
                if idx == -1:
                    break
                result = result[:idx] + orig_acr + result[idx + len(sub_acr):]
                replaced += 1
                pos = idx + len(orig_acr)

    return result


# ---------------------------------------------------------------------------
# EDITING_INSTRUCTIONS: shared instruction block included in all three editing
# prompts (paragraphs, footnotes, table cells). Update here to affect all.
# Context-specific instructions (batch delimiters, URL handling) are appended
# separately in each prompt.
# ---------------------------------------------------------------------------
EDITING_INSTRUCTIONS = """Instructions:
- Correct grammar, spelling and capitalization.
- Improve the quality of the writing: replace imprecise, informal or awkward phrasing with clear, precise and formal language appropriate for an intergovernmental document.
- Do not add or remove information beyond corrections. Do not delete words, numbers or phrases.
- Do not include any parts of the style guide in your response.
- Do not add comments about text completeness.
- Maintain consistency in terminology within and across paragraphs.
- Do not remove or alter superscripted ordinal suffixes (e.g. "5th", "3rd" — leave "th" and "rd" exactly as they appear).
- The first word after an introductory phrase and a colon should be lower-cased unless it is a proper noun or upper-cased in the terminology file (e.g. "Adaptation Fund: at SBI 64..." not "At SBI 64...").
- ACRONYMS: leave all acronyms exactly as they appear — do not expand, define, or write out the full term.
- NUMBERS: numbers 9 and below must be written as words (one, two, ... nine); numbers 10 and above as digits. Apply this to every number in the text without exception (e.g. "fifty-six" → "56", "ten" → "10", "three" → "three" [correct, keep], "11" → "11" [correct, keep], "six" → "six" [correct, keep — it is 9 or below]).
- Do not edit URLs or add any comments about them. Leave URLs exactly as they appear.
- LIST PUNCTUATION: within a sentence containing a list of items, use consistent separators throughout. If the first separator between items is a semicolon, all subsequent separators must also be semicolons. If the first separator is a comma, all must be commas.
- "Party" and "Parties" are always capitalised when referring to Parties to the UNFCCC, the Kyoto Protocol or the Paris Agreement — including in compound modifiers such as "Party-driven", "Party-led", "non-Party". Do not lowercase these.

SPELLING — apply United Nations spelling throughout (this is mandatory, check every word):
- Use -ize not -ise: maximize, organize, prioritize, recognize, finalize, utilize, emphasize, mobilize, stabilize, operationalize, revitalize, familiarize, etc.
- Use -ization not -isation: organization, recognition, utilization, operationalization, etc.
- Use -yse not -yze: analyse (not analyze), paralyse (not paralyze)
- Check every -ise word in the text and convert to -ize. Examples: organise→organize, maximise→maximize, recognise→recognize, finalise→finalize, utilise→utilize, emphasise→emphasize, mobilise→mobilize, revitalise→revitalize, familiarise→familiarize, prioritise→prioritize, stabilise→stabilize"""

# ---------------------------------------------------------------------------
# All helper functions below are UNCHANGED from v1.
# ---------------------------------------------------------------------------

def create_llm_bot(metadata: dict) -> InterfaceToLLM:
    client = AnthropicClient(default_model=metadata['model'])
    interface = InterfaceToLLM(client=client, system_prompt=metadata["system_prompt"])
    interface.authenticate(api_key=metadata["api_key"])
    return interface

def ensure_folder_exists(folder_path):
    if not os.path.exists(folder_path):
        os.makedirs(folder_path, exist_ok=True)
    if not os.access(folder_path, os.W_OK):
        try:
            import stat
            os.chmod(folder_path, stat.S_IWRITE)
        except Exception as e:
            print(f"Warning: Could not modify folder permissions for {folder_path} -> {e}")

# Regex to detect URLs — used to guard against model commentary about them
_URL_PATTERN = re.compile(r'https?://\S+')

def strip_url_comments(edited_text: str, original_text: str) -> str:
    """
    Restores any URLs the model altered, and removes any commentary the model
    appended *about* URLs — but only if that commentary was not in the original.
    """
    # Step 1: restore any URLs the model may have changed
    original_urls = _URL_PATTERN.findall(original_text)
    edited_urls = _URL_PATTERN.findall(edited_text)
    if original_urls and original_urls != edited_urls:
        for orig_url, edit_url in zip(original_urls, edited_urls):
            if orig_url != edit_url:
                edited_text = edited_text.replace(edit_url, orig_url, 1)

    # Step 2: remove URL commentary the model appended (e.g. "[URL not edited]").
    # Only strip if: edited is longer than original AND ends with a commentary phrase
    # that was NOT in the original. Avoids stripping legitimate URL-containing text.
    if len(edited_text) > len(original_text) + 10:
        url_comment_pattern = re.compile(
            r'\s*[\(\[]?[^.!?\n]{0,80}(?:URL|has not been edited|preserved as.is)[^.!?\n]{0,80}[.!?\)\]]*\s*$',
            re.IGNORECASE
        )
        cleaned = url_comment_pattern.sub('', edited_text).strip()
        if cleaned and len(cleaned) >= len(original_text) * 0.7:
            return cleaned

    return edited_text


def edit_identification(original_paragraph, edited_paragraph):
    dmp = diff_match_patch()
    diffs = dmp.diff_main(original_paragraph, edited_paragraph)
    dmp.diff_cleanupSemantic(diffs)
    return diffs

def add_insertion(paragraph, text, trackchanges_id, rsid):
    run = paragraph.add_run(text)
    ins_elem = OxmlElement('w:ins')
    ins_elem.set(qn('w:author'), "UNFCCC_ProofReader")
    ins_elem.set(qn('w:date'), datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'))
    ins_elem.set(qn('w:id'), str(trackchanges_id))
    run._element.set(qn('w:rsidR'), rsid)
    ins_elem.append(run._element)
    paragraph._element.append(ins_elem)

def add_deletion(paragraph, text, trackchanges_id, rsid):
    del_elem = OxmlElement('w:del')
    del_elem.set(qn('w:author'), "UNFCCC_ProofReader")
    del_elem.set(qn('w:date'), datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'))
    del_elem.set(qn('w:id'), str(trackchanges_id))
    run_elem = OxmlElement('w:r')
    run_elem.set(qn('w:rsidDel'), rsid)
    deltext_elem = OxmlElement('w:delText')
    deltext_elem.set(qn('xml:space'), "preserve")
    deltext_elem.text = text
    run_elem.append(deltext_elem)
    del_elem.append(run_elem)
    paragraph._element.append(del_elem)

def generateRsid(num):
    if num > 16**9:
        raise("rsid int too large (>=16**9)")
    n = 1
    while num//(16**n) != 0: n += 1
    rsid = "0"*(8-n) + str(hex(num))[2:]
    return rsid

def build_char_rpr_map(paragraph):
    """
    Builds a list of (rpr_copy_or_None) for every character in paragraph.text,
    so each character knows which run properties it came from.
    This lets us preserve bold/italic/size on a per-character basis when
    rebuilding the paragraph from diff edits.
    """
    import copy
    char_rprs = []
    for run in paragraph.runs:
        rpr = run._element.find(qn('w:rPr'))
        rpr_copy = copy.deepcopy(rpr) if rpr is not None else None
        for _ in run.text:
            char_rprs.append(rpr_copy)
    return char_rprs


def get_dominant_rpr(paragraph):
    """
    Returns the most common w:rPr in the paragraph (by character count),
    used as a fallback for inserted text where we have no original character to map to.
    """
    import copy
    from collections import Counter
    rpr_counts = Counter()
    rpr_map = {}
    for run in paragraph.runs:
        rpr = run._element.find(qn('w:rPr'))
        key = LET.tostring(rpr).decode() if rpr is not None else "__none__"
        rpr_counts[key] += len(run.text)
        if key not in rpr_map:
            rpr_map[key] = copy.deepcopy(rpr) if rpr is not None else None
    if not rpr_counts:
        return None
    dominant_key = rpr_counts.most_common(1)[0][0]
    return rpr_map[dominant_key]


def make_run_with_rpr(text, rpr_elem, del_text=False):
    """
    Creates a w:r with the given text and a copy of rpr_elem (if any).
    If del_text=True, uses w:delText instead of w:t.
    """
    import copy
    run = OxmlElement('w:r')
    if rpr_elem is not None:
        run.append(copy.deepcopy(rpr_elem))
    tag = 'w:delText' if del_text else 'w:t'
    t = OxmlElement(tag)
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    run.append(t)
    return run


def rpr_key(rpr_elem):
    """Returns a hashable key for comparing two rPr elements."""
    if rpr_elem is None:
        return None
    return LET.tostring(rpr_elem).decode()


def split_segment_by_rpr(text, start_pos, char_rprs, fallback_rpr):
    """
    Splits a diff segment into sub-segments wherever the rPr changes,
    so each sub-segment becomes a correctly-formatted run.
    Returns list of (sub_text, rpr_elem).
    """
    if not text:
        return []
    segments = []
    current_text = text[0]
    current_rpr = char_rprs[start_pos] if start_pos < len(char_rprs) else fallback_rpr
    for i, ch in enumerate(text[1:], 1):
        pos = start_pos + i
        rpr = char_rprs[pos] if pos < len(char_rprs) else fallback_rpr
        if rpr_key(rpr) != rpr_key(current_rpr):
            segments.append((current_text, current_rpr))
            current_text = ch
            current_rpr = rpr
        else:
            current_text += ch
    segments.append((current_text, current_rpr))
    return segments


def is_superscript_run(run_elem):
    """Returns True if this run has vertAlign superscript or subscript."""
    ns_uri = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    rpr = run_elem.find(f"{{{ns_uri}}}rPr")
    if rpr is None:
        return False
    vert = rpr.find(f"{{{ns_uri}}}vertAlign")
    if vert is None:
        return False
    return vert.get(f"{{{ns_uri}}}val", "") in ("superscript", "subscript")


def build_para_info(paragraph):
    """
    Builds everything needed to rebuild a paragraph after editing:
    - char_rprs: rPr per character of paragraph.text
    - super_positions: list of (char_pos_after, run_element) for superscript runs,
      so they can be reinserted after the character they follow
    - fallback_rpr: most common non-super rPr (for inserted text)

    Uses paragraph.text as the canonical text, matching what the model sees.
    """
    import copy
    from collections import Counter
    ns_uri = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    char_rprs = []
    super_positions = []  # (char_pos_after, run_element)
    rpr_counts = Counter()
    rpr_store = {}
    pos = 0  # position in paragraph.text

    for run in paragraph.runs:
        rpr = run._element.find(f"{{{ns_uri}}}rPr")
        rpr_copy = copy.deepcopy(rpr) if rpr is not None else None
        run_text = run.text or ""

        if is_superscript_run(run._element):
            # Superscript run: record its position (after pos), don't add to char_rprs
            # The text of this run IS in paragraph.text, so we must account for it
            super_positions.append((pos, run._element, len(run_text)))
            # We still add chars to char_rprs so positions stay aligned with paragraph.text
            for _ in run_text:
                char_rprs.append(rpr_copy)
            pos += len(run_text)
        else:
            key = rpr_key(rpr_copy)
            rpr_counts[key] += max(len(run_text), 1)
            if key not in rpr_store:
                rpr_store[key] = rpr_copy
            for _ in run_text:
                char_rprs.append(rpr_copy)
            pos += len(run_text)

    fallback_rpr = rpr_store[rpr_counts.most_common(1)[0][0]] if rpr_counts else None
    return char_rprs, super_positions, fallback_rpr


def insert_paragraph_revision(original_paragraph, edits, trackchanges_id, rsid_num):
    """
    Rebuilds a paragraph from diff edits, preserving per-character formatting
    and reinserting superscript runs at their correct positions.

    The diff is between paragraph.text (what the model saw) and the edited text,
    so positions align with char_rprs and super_positions.
    """
    import copy

    char_rprs, super_positions, fallback_rpr = build_para_info(original_paragraph)

    p_elem = original_paragraph._p
    for child in list(p_elem):
        if child.tag != qn('w:pPr'):
            p_elem.remove(child)

    orig_pos = 0  # position in original paragraph.text

    def flush_supers_up_to(target_pos, is_deletion=False):
        """Reinsert any superscript runs whose start position <= target_pos."""
        nonlocal super_positions
        remaining = []
        for super_start, super_elem, super_len in super_positions:
            if super_start <= target_pos:
                if not is_deletion:
                    p_elem.append(copy.deepcopy(super_elem))
                # If in a deletion region, we still preserve the super run
                # (don't delete ordinal suffixes)
                else:
                    p_elem.append(copy.deepcopy(super_elem))
            else:
                remaining.append((super_start, super_elem, super_len))
        super_positions[:] = remaining

    def is_super_rpr(rpr_elem):
        """Check if an rPr element has superscript/subscript vertAlign."""
        if rpr_elem is None:
            return False
        ns_uri = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        vert = rpr_elem.find(f"{{{ns_uri}}}vertAlign")
        if vert is None:
            return False
        return vert.get(f"{{{ns_uri}}}val", "") in ("superscript", "subscript")

    for operation, edit in edits:
        if operation == 0:
            for sub_text, rpr in split_segment_by_rpr(edit, orig_pos, char_rprs, fallback_rpr):
                if is_super_rpr(rpr):
                    # These characters belong to a superscript run.
                    # Flush the original run element — don't create a new run from the diff.
                    flush_supers_up_to(orig_pos + len(sub_text) - 1)
                else:
                    flush_supers_up_to(orig_pos + len(sub_text) - 1)
                    p_elem.append(make_run_with_rpr(sub_text, rpr))
                orig_pos += len(sub_text)

        elif operation == -1:
            del_elem = OxmlElement('w:del')
            del_elem.set(qn('w:author'), "UNFCCC_ProofReader")
            del_elem.set(qn('w:date'), datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'))
            del_elem.set(qn('w:id'), str(trackchanges_id))
            has_del_content = False
            for sub_text, rpr in split_segment_by_rpr(edit, orig_pos, char_rprs, fallback_rpr):
                if is_super_rpr(rpr):
                    # Superscript in deleted region — preserve it, don't mark as deleted
                    if has_del_content:
                        p_elem.append(del_elem)
                        trackchanges_id += 1
                        del_elem = OxmlElement('w:del')
                        del_elem.set(qn('w:author'), "UNFCCC_ProofReader")
                        del_elem.set(qn('w:date'), datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'))
                        del_elem.set(qn('w:id'), str(trackchanges_id))
                        has_del_content = False
                    flush_supers_up_to(orig_pos + len(sub_text) - 1, is_deletion=True)
                else:
                    del_elem.append(make_run_with_rpr(sub_text, rpr, del_text=True))
                    has_del_content = True
                orig_pos += len(sub_text)
            if has_del_content:
                p_elem.append(del_elem)
                trackchanges_id += 1

        elif operation == 1:
            if orig_pos > 0 and orig_pos - 1 < len(char_rprs):
                insert_rpr = char_rprs[orig_pos - 1]
            else:
                insert_rpr = fallback_rpr
            ins_elem = OxmlElement('w:ins')
            ins_elem.set(qn('w:author'), "UNFCCC_ProofReader")
            ins_elem.set(qn('w:date'), datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'))
            ins_elem.set(qn('w:id'), str(trackchanges_id))
            ins_elem.append(make_run_with_rpr(edit, insert_rpr))
            p_elem.append(ins_elem)
            trackchanges_id += 1
            rsid_num += 1

    # Flush any remaining superscript runs at the end
    for _, super_elem, _ in super_positions:
        p_elem.append(copy.deepcopy(super_elem))

    return trackchanges_id, rsid_num

def extract_spelling_capitalization_changes(edits):
    changes = []
    i = 0
    while i < len(edits) - 1:
        if edits[i][0] == -1 and edits[i+1][0] == 1:
            old_text = edits[i][1]
            new_text = edits[i+1][1]
            if old_text.lower() == new_text.lower():
                changes.append((old_text, new_text))
            i += 2
        else:
            i += 1
    return changes

# Load terminology
with open(os.path.join(os.path.dirname(__file__), "terminology.md"), "r", encoding="utf-8") as f:
    terminology_data = f.read()

with open(os.path.join(os.path.dirname(__file__), "country_names.md"), "r", encoding="utf-8") as f:
    country_names_data = f.read()

def get_relevant_terms(paragraph_text, terminology_data):
    relevant_terms = []
    for line in terminology_data.splitlines():
        if not line.strip() or line.isupper():  # skip headers
            continue
        # Extract just the primary term (before '<', '(', '=' or explanatory text)
        primary = re.split(r'[<(=\[]', line)[0].strip().rstrip(':')
        if len(primary) < 4:  # skip very short tokens prone to false matches
            continue
        if primary.lower() in paragraph_text.lower():
            relevant_terms.append(line.strip())
    return "\n".join(relevant_terms[:20])
# ---------------------------------------------------------------------------
# System prompt — UNCHANGED from v1 (abbreviated here, paste your full one)
# ---------------------------------------------------------------------------
system_prompt = """
This is the UNFCCC Editorial Style Guide, to be followed at all times:
    
SESSION NUMBERING AND PREPOSITIONS  
- When the name of a body is written out, the session number should be written out (e.g., "fiftieth session of the Subsidiary Body for Implementation").
- When an acronym is used, the session number should be in digits, with a space before it (e.g., "SB 58", "CMP 13", "CMA 7", "COP 28").
- Use "at" before session numbers, not "by" (e.g., "done at CMA 5", not "done by CMA 5").

MEETING NUMBERING
- Unlike with session numbers, even when the word "meeting" appears, meeting numbers should be in ordinal digits (e.g. "29th meeting of the TEC", "3rd meeting of the PCCB"). Do not alter ordinal suffixes ("th", "rd", "st", "nd") — their formatting is handled separately.
- If the word "meeting" is omitted, meeting numbers can be written like session numbers (e.g. "TEC 29", "AC 17").

OTHER NUMBERS
- Numbers 9 and lower are written out (write e.g. "eight", not "8"); numbers 10 and higher are given as digits (write e.g. "16", not "sixteen"). If numbers are given for a series of comparable items, and one of the 
numbers is 10 or larger, all the numbers in that series should be given as digits (e.g. "24 fixed-term staff, 12 temporary staff and 8 consultants").
- Ordinal number suffixes (e.g. "th" in "5th") are already superscripted in the document. Do not add, remove or alter superscript formatting on ordinal suffixes.

ABBREVIATIONS AND ACRONYMS
- Leave every acronym exactly as it appears. Do NOT expand, define or write out any acronym — not even on first mention. If "NDCs" appears, leave it as "NDCs". If "CTCN" appears, leave it as "CTCN". No exceptions.
- Do NOT substitute one acronym for another. CTC and CTCN are distinct — never change one to the other.
- Avoid possessive forms of acronyms (e.g., "SCF's" → "the SCF co-chairs").
- "United Nations" should only be abbreviated when part of an official acronym.
- Acronyms should not be preceded by "the". Exceptions (sentences only, not lists):
    - UNFCCC, COP, CMP, CMA, SBI, SBSTA, ES, AC, TEC, CTC, CTCN, LEG, SCF, CDM, GCF, GEF, LDCF, NWP and IPCC
SPELLING  
- Use United Nations spelling, which follows British spelling, except for "-ize" and "-yse" words. For example:  
    - "organise" → "organize" (use "-ize" endings)  
    - "analyze" → "analyse", "paralyze" → "paralyse" (use "-yse" endings)  
    - "organisation" → "organization"
    - ✅ generalization
- Use the correct spelling for all terms in the terminology file.

CAPITALIZATION
    - Names of pathways should be lower-cased.
    - Headings should be in sentence case, except for terms that are in upper case in the terminology file.
    - The first word after an introductory phrase and a colon should lower-cased unless a proper noun or upper-cased in the terminology file.
    - "Annex" (capital A) is used only when referring to an annex to the Convention or the Kyoto Protocol (e.g. "Annex I", "Annex II"). 
    In all other cases use "annex" (lower case), including annexes to decisions, reports or any other documents.

OXFORD COMMA  
- **Do not use** the Oxford comma unless necessary for clarity.  
    - ❌ "organs, organizations, and bodies"  
    - ✅ "organs, organizations and bodies"  
    - ❌ "adjusts, reviews, and submits"  
    - ✅ "adjusts, reviews and submits"  
- Use a **serial comma** in complex lists to avoid ambiguity:  
    - ✅ "...education and literacy programmes, and health and social support programmes"

HYPHENATION  
- Two-word modifiers **must** be hyphenated (e.g., "long-term plan", "climate-dependent pathways").  
- Three or more words modifying a noun should **not** be hyphenated unless a pre-hyphenated term is included (e.g., "capacity-building-related investments").

QUOTES
- Quoted terms should be inside punctuation marks that immediately follow them.
    - ❌ referred to as 'business as usual,'
    - ✅ referred to as 'business as usual',  

LAYOUT
Subparagraphs and bullets start with a capital letter, end with a semicolon; no "and" before the final item. The final item ends with a full stop.

CURRENCIES  
- Outside tables, use **currency abbreviations**, not symbols (e.g., "USD 100 million", not "$100 million").  
- **Use currency symbols ($, €, £) in tables only.**

SPECIAL CASE WORDS  
- "interlinkage" can be used in **singular** or **plural**.  
- Use "at the local level", **not** "at local levels".

NAMES OF PERSONS
- Names should not be preceded by prefixes such "Mr.", "Mr", "Ms." and "Ms"
    - ❌ Mr. Simon Stiell
    - ✅ Simon Stiell
- Titles like "Her Excellency", "Ambassador" and "Dr" should be omitted. Delete them if they appear in the text.

SYMBOLS
- There should be a space between the Celsius symbol and the number:
    - ❌ 1.5°C 
    - ✅ 1.5 °C 

MISCELLANEOUS RULES 
- Use **"i.e."** and **"e.g."** only inside parentheses. Do **not** follow them with a comma.  
    - ✅ (e.g. renewable energy, energy efficiency)  
    - ❌ (e.g., renewable energy, energy efficiency)  
- Outside parentheses:  
    - Replace **"i.e."** with **"that is"**  
    - Replace **"e.g."** with **"such as"** or **"for example"**  
    - ✅ "Countries, such as Germany and France, have adopted this policy."  
    - ❌ "Countries, e.g., Germany and France, have adopted this policy."  
- Use a **comma before "etc."** in running text. If the sentence continues after "etc.," follow it with another comma.  
    - ✅ "Countries invest in renewable energy, efficiency, grid modernization, etc., to meet targets."  
    - ❌ "Countries invest in renewable energy, efficiency, grid modernization etc."  
- **Treat "Parties" as things, not people.** Use **"Parties that"**, not **"Parties who"**.  
    - ✅ "Parties that signed the agreement..."  
    - ❌ "Parties who signed the agreement..."  
- Write **"under and outside the Convention"**, not **"under and outside of the Convention"**.
- When quoting the text, place the closing parenthesis inside the punctation (i.e. inside the full stop or the comma).
- Do not replace curly quotes with straight quotes or curley apostrophes with straight apostrophes.
- "ES" refers to "Executive Secretary".
"""

# ---------------------------------------------------------------------------
# CHANGE: Pre-pass — extract a document-wide consistency resolution table
# before any editing begins.
#
# This replaces the growing consistency_log approach. Instead of recording
# what the model happened to change paragraph-by-paragraph, we ask the model
# upfront to resolve every ambiguous term/spelling/capitalization across the
# whole document against the style guide, producing a lookup table that gets
# injected into every subsequent editing prompt.
#
# The result is a dict like:
#   { "Organisation": "Organization", "COP26": "COP 26", ... }
# ---------------------------------------------------------------------------

def build_consistency_table(full_document_text: str, country_names_data: str, terminology_data: str, interface: InterfaceToLLM) -> dict:
    prompt = f"""
You are reviewing a full document before editing begins. Your task is NOT to edit the text yet.

Scan the entire document and identify every term, acronym, acronym definition, spelling variant or capitalization issue
that conflicts with the UNFCCC Editorial Style Guide (in your system prompt) and the terms in the terminology file.

Pay attention the rules in the Style Guide, but also to spelling and capitalization of the standard terms and definitions of acronyms in the
terminology list. Do not assume these terms are definitions are proper nouns; follow the terminology file.

Examples:
- Terminology shows "NDC = nationally determined contribution" → flag "Nationally Determined Contribution" or "Nationally determined contribution" → "nationally determined contribution"
- Terminology shows "GST = global stocktake" → flag "Global Stocktake" or "Global stocktake" → "global stocktake"
- Terminology shows "NAP = Trust Fund for Participation in the UNFCCC Process" → flag "Trust Fund for Participation in the UNFCCC Process" → "trust fund for participation in the UNFCCC process"
- Terminology shows "technology mechanism" → flag "technology mechanism" → "Technology Mechanism"

Apply this to ALL standard terms and acronym definitions in the terminology file, regardless of whether the term appears in singular or plural form and whether it's capitalized or lower-cased.

If an acronym in the document appears in that form in the terminology file, do not flag it for change to another acronym 
(e.g. don't flag "NDE" for correction to "NDC"). Do not replace acronym definitions with acronyms (don't abbreviate terms).

For "Annex"/"annex": flag "Annex" (capital) only when it does NOT refer to an annex of the Convention or the Kyoto Protocol — in those cases it should remain "Annex". Flag "Annex" as needing correction to "annex" when it refers to annexes to decisions, reports or other documents.

Also apply the following country name rules: enforce short-form country names and correct
any spelling errors using the reference list below.

Produce a JSON object where each key is a non-standard form found in the document and each
value is the correct form. Output raw JSON only.

Country name reference:
{country_names_data}

Terminology:
{terminology_data}

Document:
{full_document_text}
"""
    # ... rest unchanged
    response = interface.get_response(prompt=prompt)
    raw = response["text"].strip()

    # Strip accidental markdown fences if the model adds them despite instructions
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
        raw = raw.strip()

    try:
        table = json.loads(raw)
        print(f"[Pre-pass] Consistency table built: {len(table)} entries.")
        return table
    except json.JSONDecodeError as e:
        print(f"[Pre-pass] Warning: could not parse consistency table JSON: {e}")
        print(f"[Pre-pass] Raw response was: {raw[:500]}")
        return {}


def format_consistency_table(table: dict) -> str:
    """
    Formats the consistency table as a compact instruction block to inject
    into every editing prompt.
    """
    if not table:
        return ""
    lines = ["Document-wide consistency resolutions (apply these uniformly throughout):"]
    for wrong, correct in table.items():
        lines.append(f'  - "{wrong}" → "{correct}"')
    return "\n".join(lines) + "\n\n"


# ---------------------------------------------------------------------------
# CHANGE: Section-level batching for paragraph editing.
#
# Instead of one LLM call per paragraph, we group paragraphs into batches
# and send them together. The model sees local context (surrounding sentences)
# which helps it make consistent decisions within a section.
#
# Each paragraph is delimited with a unique marker so we can split the
# response back into individual edited paragraphs reliably.
# ---------------------------------------------------------------------------

PARA_DELIMITER = "|||PARA_{index}|||"

def build_batch_prompt(paragraphs_with_text: list[tuple], consistency_block: str) -> str:
    """
    paragraphs_with_text: list of (index, original_text) tuples for this batch.
    Returns a prompt asking the model to edit all paragraphs and return them
    in the same delimited format.
    """
    # Collect relevant terms across all paragraphs in the batch
    if USE_RELEVANT_TERMS:
        combined_text = " ".join(text for _, text in paragraphs_with_text)
        relevant = get_relevant_terms(combined_text, terminology_data)
        terms_text = f"\n\nRelevant terminology:\n{relevant}" if relevant else ""
    else:
        terms_text = ""

    # Build delimited input block
    input_block = ""
    for index, text in paragraphs_with_text:
        input_block += f"{PARA_DELIMITER.format(index=index)}\n{text}\n"

    prompt = f"""Edit each paragraph below according to the UNFCCC Editorial Style Guide (in your system prompt).

{EDITING_INSTRUCTIONS}
- Return ONLY the edited paragraphs, each preceded by its delimiter exactly as shown.
- Do not add any text outside the delimited paragraphs.

{consistency_block}{terms_text}

Paragraphs to edit:

{input_block}"""
    return prompt


def parse_batch_response(response_text: str, indices: list[int]) -> dict[int, str]:
    """
    Splits the model's batch response back into per-paragraph texts.
    Returns a dict { original_index: edited_text }.
    Falls back gracefully if parsing fails for any paragraph.
    """
    result = {}
    for i, index in enumerate(indices):
        delimiter = PARA_DELIMITER.format(index=index)
        next_delimiter = PARA_DELIMITER.format(index=indices[i+1]) if i+1 < len(indices) else None

        start = response_text.find(delimiter)
        if start == -1:
            print(f"[Batch parse] Warning: delimiter not found for paragraph {index}. Keeping original.")
            result[index] = None  # Signal to keep original
            continue

        start += len(delimiter)
        if next_delimiter:
            end = response_text.find(next_delimiter)
            if end == -1:
                end = len(response_text)
        else:
            end = len(response_text)

        parsed = response_text[start:end].strip()
        # Guard: if the model accidentally included the next delimiter inside
        # the paragraph text, strip everything from the delimiter onwards.
        for other_index in indices:
            stray = PARA_DELIMITER.format(index=other_index)
            if stray in parsed:
                parsed = parsed[:parsed.index(stray)].strip()
                break
        result[index] = parsed

    return result

def edit_footnote(footnote_paragraph, interface, relevant_terms, consistency_block, system_prompt, trackchanges_id, rsid_num):
    """
    Edits plain-text runs in a footnote paragraph while preserving original child order.
    Strategy: snapshot all children, classify each as pPr/marker/hyperlink/run,
    send only run text to model, rebuild only the runs in their original position.
    All other children (pPr, hyperlinks, marker) are preserved in place.
    """
    import copy
    ns_uri = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    xml_ns = "http://www.w3.org/XML/1998/namespace"
    ns_dict = {"w": ns_uri}

    # Snapshot and classify all children
    children_info = []  # (element, category)
    for child in list(footnote_paragraph):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'pPr':
            children_info.append((child, 'pPr'))
        elif tag == 'hyperlink':
            children_info.append((child, 'hyperlink'))
        elif tag == 'r':
            is_marker = (
                child.find(".//w:footnoteReference", namespaces=ns_dict) is not None or
                child.find(".//w:footnoteRef", namespaces=ns_dict) is not None
            )
            children_info.append((child, 'marker' if is_marker else 'run'))
        else:
            children_info.append((child, 'other'))

    # Extract plain text from run children only
    original_text = ''.join(
        ''.join(child.itertext()) for child, cat in children_info if cat == 'run'
    ).strip()

    if not original_text or re.match(r'^https?://\S+$', original_text):
        return trackchanges_id, rsid_num  # nothing editable

    # Capture run formatting from first non-empty run
    fn_rpr = None
    for child, cat in children_info:
        if cat == 'run':
            rpr = child.find(f"{{{ns_uri}}}rPr")
            if rpr is not None:
                fn_rpr = copy.deepcopy(rpr)
                break

    # Send plain text to model
    terms_text = f"\n\nRelevant terminology:\n{relevant_terms}" if relevant_terms else ""
    response = interface.get_response(prompt=f"""
Edit the text according to the UNFCCC Editorial Style Guide (in your system prompt).

{EDITING_INSTRUCTIONS}

{consistency_block}{terms_text}

Text to edit:

{original_text}
""")
    edited_text = protect_acronyms(
        apply_forced_corrections(strip_url_comments(response["text"].strip(), original_text)),
        original_text
    )

    dmp = diff_match_patch()
    diffs = dmp.diff_main(original_text, edited_text)
    dmp.diff_cleanupSemantic(diffs)

    # Build new run elements from diff
    def _fn_run(text_content, del_text=False):
        r = LET.Element(f"{{{ns_uri}}}r")
        if fn_rpr is not None:
            r.append(copy.deepcopy(fn_rpr))
        t_tag = f"{{{ns_uri}}}delText" if del_text else f"{{{ns_uri}}}t"
        t = LET.Element(t_tag)
        t.set(f"{{{xml_ns}}}space", "preserve")
        t.text = text_content
        r.append(t)
        return r

    new_run_elements = []
    for op, text in diffs:
        if op == 0:
            new_run_elements.append(_fn_run(text))
        elif op == 1:
            ins = LET.Element(f"{{{ns_uri}}}ins")
            ins.set(f"{{{ns_uri}}}author", "UNFCCC_ProofReader")
            ins.set(f"{{{ns_uri}}}date", datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'))
            ins.set(f"{{{ns_uri}}}id", str(trackchanges_id))
            ins.append(_fn_run(text))
            new_run_elements.append(ins)
            trackchanges_id += 1
            rsid_num += 1
        elif op == -1:
            deletion = LET.Element(f"{{{ns_uri}}}del")
            deletion.set(f"{{{ns_uri}}}author", "UNFCCC_ProofReader")
            deletion.set(f"{{{ns_uri}}}date", datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'))
            deletion.set(f"{{{ns_uri}}}id", str(trackchanges_id))
            deletion.append(_fn_run(text, del_text=True))
            new_run_elements.append(deletion)
            trackchanges_id += 1

    # Rebuild paragraph in original child order:
    # remove all children, reinsert each in order,
    # replacing ALL run children with new_run_elements at the position of the first run.
    for child, _ in children_info:
        footnote_paragraph.remove(child)

    first_run_inserted = False
    for child, cat in children_info:
        if cat == 'run':
            if not first_run_inserted:
                for new_elem in new_run_elements:
                    footnote_paragraph.append(new_elem)
                first_run_inserted = True
            # skip subsequent original run children — replaced by new_run_elements
        else:
            footnote_paragraph.append(child)

    return trackchanges_id, rsid_num


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def editorial_proofreader(docx_file_name: str, styles_excluded: list[str], use_both_models: bool = False):
    today_folder = f"data/input/{datetime.now().strftime('%d-%m-%Y')}/"
    ensure_folder_exists(today_folder)

    file_path = os.path.join(today_folder, f"{docx_file_name}.docx")
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Error: The file {file_path} does not exist!")

    doc = docx.Document(file_path)
    load_dotenv()

    # Pre-pass always uses Sonnet for maximum accuracy on document-wide analysis.
    # Paragraph/table/footnote editing uses Haiku (single-model) or Haiku→Sonnet (two-model).
    interface_sonnet = create_llm_bot(metadata={
        "model": "claude-sonnet-4-6",
        "system_prompt": system_prompt,
        "api_key": os.getenv("ANTHROPIC_API_KEY")
    })
    interface_haiku = create_llm_bot(metadata={
        "model": "claude-haiku-4-5-20251001",
        "system_prompt": system_prompt,
        "api_key": os.getenv("ANTHROPIC_API_KEY")
    })

    # -----------------------------------------------------------------------
    # CHANGE: Pre-pass — build consistency table from the full document text.
    #
    # We extract all paragraph text, join it, and send it in one shot.
    # For very long documents (>100k tokens) you may want to chunk this, but
    # for typical UNFCCC documents (up to ~50 pages) it fits comfortably.
    # -----------------------------------------------------------------------
    print("[Pre-pass] Extracting full document text for consistency analysis...")
    full_doc_paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    full_document_text = "\n\n".join(full_doc_paragraphs)

    consistency_table = build_consistency_table(full_document_text, country_names_data, terminology_data, interface_sonnet)
    consistency_block = format_consistency_table(consistency_table)
    if consistency_table:
        print("[Pre-pass] Consistency resolutions:")
        for wrong, correct in consistency_table.items():
            print(f"  '{wrong}' → '{correct}'")
    else:
        print("[Pre-pass] No consistency issues found.")

    # -----------------------------------------------------------------------
    # Edit paragraphs — CHANGE: batched instead of one call per paragraph.
    #
    # We collect eligible paragraphs (non-empty, not in excluded styles),
    # group them into batches of BATCH_SIZE, send each batch as a single
    # LLM call, then parse the response and apply tracked changes.
    # -----------------------------------------------------------------------
    trackchanges_id = 1
    rsid_num = 1
    consistency_log = ""  # Kept for the log file; no longer fed back to the model

    # Collect eligible paragraphs with their positions
    eligible = [
        (i, p) for i, p in enumerate(doc.paragraphs)
        if p.text.strip() and p.style.name not in styles_excluded
    ]

    print(f"[Editing] {len(eligible)} eligible paragraphs in {(len(eligible) + BATCH_SIZE - 1) // BATCH_SIZE} batches.")

    for batch_start in range(0, len(eligible), BATCH_SIZE):
        batch = eligible[batch_start: batch_start + BATCH_SIZE]
        batch_indices = [i for i, _ in batch]
        batch_texts = [(i, p.text) for i, p in batch]

        # Log and print relevant terms for this batch
        if USE_RELEVANT_TERMS:
            combined_text = " ".join(text for _, text in batch_texts)
            batch_terms = get_relevant_terms(combined_text, terminology_data)
            if batch_terms:
                batch_num = batch_start // BATCH_SIZE + 1
                print(f"  [Batch {batch_num}] Relevant terms: {batch_terms.replace(chr(10), ' | ')}")
                consistency_log += f"\nBatch {batch_num} relevant terms:\n{batch_terms}\n"

        if use_both_models:
            # --- Two-model path: Haiku first pass (fast, cheap), Sonnet second pass (thorough) ---
            interface_mini = create_llm_bot(metadata={
                "model": "claude-haiku-4-5-20251001",
                "system_prompt": system_prompt,
                "api_key": os.getenv("ANTHROPIC_API_KEY")
            })
            prompt_mini = build_batch_prompt(batch_texts, consistency_block)
            response_mini = interface_mini.get_response(prompt=prompt_mini)
            # Parse mini response into per-paragraph texts for the second pass
            mini_results = parse_batch_response(response_mini["text"], batch_indices)

            # Build second-pass batch using Haiku-edited text where available
            second_pass_texts = [
                (i, mini_results.get(i) or p.text)
                for i, p in batch
            ]
            interface_sonnet = create_llm_bot(metadata={
                "model": "claude-sonnet-4-6",
                "system_prompt": system_prompt,
                "api_key": os.getenv("ANTHROPIC_API_KEY")
            })
            prompt_sonnet = build_batch_prompt(second_pass_texts, consistency_block)
            response_sonnet = interface_sonnet.get_response(prompt=prompt_sonnet)
            edited_results = parse_batch_response(response_sonnet["text"], batch_indices)

        else:
            # --- Single-model path: Haiku only ---
            prompt = build_batch_prompt(batch_texts, consistency_block)
            response = interface_haiku.get_response(prompt=prompt)
            edited_results = parse_batch_response(response["text"], batch_indices)

        # Apply diffs for each paragraph in this batch
        for i, paragraph in batch:
            edited_text = edited_results.get(i)
            if edited_text is None:
                # Parsing failed for this paragraph — skip silently, keep original
                print(f"  [Warning] Skipping paragraph {i} (parse failed), keeping original.")
                continue

            edited_text = protect_acronyms(apply_forced_corrections(strip_url_comments(edited_text, paragraph.text)), paragraph.text)

            # Safety guard: if the model has drastically shortened the paragraph
            # (less than 60% of original length), it has likely hallucinated a
            # truncation. Keep the original and log a warning.
            if len(paragraph.text) > 50 and len(edited_text) < 0.6 * len(paragraph.text):
                print(f"  [Warning] Paragraph {i} edited text is suspiciously short "
                      f"({len(edited_text)} vs {len(paragraph.text)} chars). Keeping original.")
                consistency_log += f"WARNING: Paragraph {i} skipped (edited too short)\n"
                continue

            edits = edit_identification(original_paragraph=paragraph.text, edited_paragraph=edited_text)

            # Log spelling/capitalization changes (for the log file)
            changes = extract_spelling_capitalization_changes(edits)
            if changes:
                for old, new in changes:
                    consistency_log += f"Changed '{old}' to '{new}'\n"

            trackchanges_id, rsid_num = insert_paragraph_revision(
                original_paragraph=paragraph,
                edits=edits,
                trackchanges_id=trackchanges_id,
                rsid_num=rsid_num
            )

        print(f"  [Editing] Batch {batch_start // BATCH_SIZE + 1} done.")

    # -----------------------------------------------------------------------
    # Determine which interface to use for tables and footnotes.
    # In the two-model path the final editing model is Sonnet (interface_sonnet).
    # In the single-model path it is the main interface (also Sonnet).
    # Pre-pass always uses the main interface (Sonnet) regardless.
    # -----------------------------------------------------------------------
    # Tables and footnotes use Sonnet as the final editing model in the two-model
    # path, and Haiku in the single-model path — consistent with paragraph editing.
    editing_interface = interface_sonnet if use_both_models else interface_haiku

    # -----------------------------------------------------------------------
    # Edit tables — CHANGE: also uses consistency_block, otherwise unchanged.
    # -----------------------------------------------------------------------
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if not cell.text.strip():
                    continue
                relevant_terms = get_relevant_terms(cell.text, terminology_data)
                terms_text = f"\n\nRelevant terminology:\n{relevant_terms}" if relevant_terms else ""

                _cell_text = editing_interface.get_response(prompt=f"""
Edit the text according to the UNFCCC Editorial Style Guide (in your system prompt).

{EDITING_INSTRUCTIONS}

{consistency_block}{terms_text}

Text to edit:

{cell.text}
""")

                cell_edited = protect_acronyms(apply_forced_corrections(strip_url_comments(_cell_text["text"], cell.text)), cell.text)
                edits = edit_identification(original_paragraph=cell.text, edited_paragraph=cell_edited)
                changes = extract_spelling_capitalization_changes(edits)
                if changes:
                    for old, new in changes:
                        consistency_log += f"Changed '{old}' to '{new}'\n"

                if cell.paragraphs:
                    trackchanges_id, rsid_num = insert_paragraph_revision(
                        original_paragraph=cell.paragraphs[0],
                        edits=edits,
                        trackchanges_id=trackchanges_id,
                        rsid_num=rsid_num
                    )

    # -----------------------------------------------------------------------
    # Footnote editing and DOCX reassembly — UNCHANGED from v1
    # -----------------------------------------------------------------------
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    input_docx_path = os.path.join(today_folder, f"{docx_file_name}.docx")
    output_docx_path = os.path.join("data/output", datetime.now().strftime('%d-%m-%Y'), f"{docx_file_name}_edited.docx")
    ensure_folder_exists(os.path.dirname(output_docx_path))

    modified_doc_xml = LET.fromstring(LET.tostring(doc._element))
    modified_paragraphs = modified_doc_xml.findall(".//w:p", namespaces=ns)

    with zipfile.ZipFile(input_docx_path, "r") as zin:
        original_doc_bytes = zin.read("word/document.xml")
    original_doc_xml = LET.fromstring(original_doc_bytes)
    original_paragraphs = original_doc_xml.findall(".//w:p", namespaces=ns)

    for orig_para, mod_para in zip(original_paragraphs, modified_paragraphs):
        footnote_refs = orig_para.findall(".//w:footnoteReference", namespaces=ns)
        if footnote_refs:
            # Remove any footnote references already in the modified paragraph
            # (they may have been dropped or misplaced during editing)
            for r in list(mod_para):
                if r.find(".//w:footnoteReference", namespaces=ns) is not None:
                    mod_para.remove(r)

            # Re-insert each footnote reference in a proper superscripted run,
            # with ", " separators between adjacent references.
            ns_uri = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            for idx, ref in enumerate(footnote_refs):
                if idx > 0:
                    # Insert a ", " separator run between markers
                    sep_run = LET.Element(f"{{{ns_uri}}}r")
                    sep_rpr = LET.Element(f"{{{ns_uri}}}rPr")
                    sep_vert = LET.Element(f"{{{ns_uri}}}vertAlign")
                    sep_vert.set(f"{{{ns_uri}}}val", "superscript")
                    sep_rpr.append(sep_vert)
                    sep_run.append(sep_rpr)
                    sep_t = LET.Element(f"{{{ns_uri}}}t")
                    sep_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                    sep_t.text = ", "
                    sep_run.append(sep_t)
                    mod_para.append(sep_run)

                # Wrap the footnoteReference in a superscripted run
                ref_run = LET.Element(f"{{{ns_uri}}}r")
                ref_rpr = LET.Element(f"{{{ns_uri}}}rPr")
                vert_align = LET.Element(f"{{{ns_uri}}}vertAlign")
                vert_align.set(f"{{{ns_uri}}}val", "superscript")
                ref_rpr.append(vert_align)
                ref_run.append(ref_rpr)
                ref_run.append(LET.fromstring(LET.tostring(ref)))
                mod_para.append(ref_run)

    updated_document_xml = LET.tostring(modified_doc_xml, encoding="utf-8", xml_declaration=True)

    with zipfile.ZipFile(input_docx_path, "r") as zin:
        original_footnotes_xml = zin.read("word/footnotes.xml")

    try:
        footnotes_tree = LET.fromstring(original_footnotes_xml)
        for footnote in footnotes_tree.findall("w:footnote", namespaces=ns):
            footnote_type = footnote.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type")
            if footnote_type in ("separator", "continuationSeparator"):
                continue
            for p in footnote.findall("w:p", namespaces=ns):
                paragraph_text = "".join(p.itertext()).strip()
                if paragraph_text:
                    relevant_terms = get_relevant_terms(paragraph_text, terminology_data)
                    trackchanges_id, rsid_num = edit_footnote(
                        p, editing_interface, relevant_terms, consistency_block, system_prompt,
                        trackchanges_id, rsid_num
                    )
        updated_footnotes_xml = LET.tostring(footnotes_tree, encoding="utf-8", xml_declaration=True)
    except Exception as e:
        print(f"Error processing footnotes.xml: {e}")
        updated_footnotes_xml = original_footnotes_xml

    with zipfile.ZipFile(input_docx_path, "r") as zin:
        with zipfile.ZipFile(output_docx_path, "w") as zout:
            for item in zin.infolist():
                if item.filename == "word/document.xml":
                    zout.writestr(item.filename, updated_document_xml)
                elif item.filename == "word/footnotes.xml":
                    zout.writestr(item.filename, updated_footnotes_xml)
                else:
                    zout.writestr(item.filename, zin.read(item.filename))

    # Write log
    with open("proofreading_log.txt", "a", encoding="utf-8") as log_file:
        log_file.write(f"\n=== Run: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} ===\n")
        log_file.write(f"Consistency table ({len(consistency_table)} entries):\n")
        for wrong, correct in consistency_table.items():
            log_file.write(f"  {wrong} -> {correct}\n")
        log_file.write(f"\nSpelling/capitalization changes applied:\n{consistency_log}\n")

    print(f"[Done] Output written to: {output_docx_path}")
